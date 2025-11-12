from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy
from html.parser import HTMLParser
import re

def _replace_text_in_paragraph(paragraph, replacements):
    """Reemplaza placeholders en un párrafo."""
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)


def _html_to_docx_paragraphs(html_content, document):
    """Convierte HTML del editor rico a párrafos de Word con formato."""
    # Remover tags HTML y convertir a texto con formato
    html_content = html_content.replace('<p>', '\n').replace('</p>', '')
    html_content = html_content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    
    # Procesar listas
    html_content = re.sub(r'<ol[^>]*>', '', html_content)
    html_content = html_content.replace('</ol>', '')
    html_content = re.sub(r'<ul[^>]*>', '', html_content)
    html_content = html_content.replace('</ul>', '')
    html_content = re.sub(r'<li[^>]*>', '\n• ', html_content)
    html_content = html_content.replace('</li>', '')
    
    # Separar por párrafos
    paragraphs_text = html_content.split('\n')
    
    for para_text in paragraphs_text:
        if not para_text.strip():
            continue
            
        # Crear párrafo
        p = document.add_paragraph()
        
        # Procesar formato inline (bold, italic, etc.)
        parts = re.split(r'(<strong>|</strong>|<b>|</b>|<em>|</em>|<i>|</i>|<u>|</u>|<strike>|</strike>)', para_text)
        
        bold = False
        italic = False
        underline = False
        strike = False
        
        for part in parts:
            if part in ['<strong>', '<b>']:
                bold = True
            elif part in ['</strong>', '</b>']:
                bold = False
            elif part in ['<em>', '<i>']:
                italic = True
            elif part in ['</em>', '</i>']:
                italic = False
            elif part == '<u>':
                underline = True
            elif part == '</u>':
                underline = False
            elif part == '<strike>':
                strike = True
            elif part == '</strike>':
                strike = False
            elif part.strip():
                # Limpiar cualquier tag residual
                clean_text = re.sub(r'<[^>]+>', '', part)
                if clean_text.strip():
                    run = p.add_run(clean_text)
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline
                    run.font.strike = strike
                    # Aplicar fuente Calibri
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)


def _get_merged_cells_info(source_table):
    """Detecta qué celdas están combinadas en la tabla fuente."""
    merged_cells = []
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            tc_pr = cell._element.tcPr
            if tc_pr is not None:
                grid_span = tc_pr.find(qn('w:gridSpan'))
                if grid_span is not None:
                    span_count = int(grid_span.get(qn('w:val')))
                    if span_count > 1:
                        merged_cells.append({
                            'row': i,
                            'start_col': j,
                            'end_col': j + span_count - 1,
                            'type': 'horizontal'
                        })
    
    return merged_cells


def _copy_table_style_and_content(source_table, target_doc):
    """Copia una tabla del template al documento destino."""
    num_rows = len(source_table.rows)
    num_cols = len(source_table.columns)
    
    new_table = target_doc.add_table(rows=num_rows, cols=num_cols)
    
    # Aplicar merges
    merged_cells_info = _get_merged_cells_info(source_table)
    for merge_info in merged_cells_info:
        try:
            row_idx = merge_info['row']
            start_col = merge_info['start_col']
            end_col = merge_info['end_col']
            
            start_cell = new_table.rows[row_idx].cells[start_col]
            end_cell = new_table.rows[row_idx].cells[end_col]
            start_cell.merge(end_cell)
        except:
            pass
    
    # Copiar contenido
    cells_visited = set()
    
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            cell_key = (i, j)
            
            if cell_key in cells_visited:
                continue
            
            cells_visited.add(cell_key)
            
            try:
                target_cell = new_table.rows[i].cells[j]
            except IndexError:
                continue
            
            # Limpiar y copiar párrafos
            for paragraph in target_cell.paragraphs:
                paragraph.clear()
            
            for k, paragraph in enumerate(cell.paragraphs):
                if k == 0:
                    target_para = target_cell.paragraphs[0]
                else:
                    target_para = target_cell.add_paragraph()
                
                target_para.alignment = paragraph.alignment
                target_para.style = paragraph.style
                
                for run in paragraph.runs:
                    target_run = target_para.add_run(run.text)
                    target_run.bold = run.bold
                    target_run.italic = run.italic
                    target_run.underline = run.underline
                    if run.font.size:
                        target_run.font.size = run.font.size
                    if run.font.name:
                        target_run.font.name = run.font.name
                    if run.font.color.rgb:
                        target_run.font.color.rgb = run.font.color.rgb
            
            # Copiar propiedades de celda
            if cell._element.tcPr is not None:
                if target_cell._element.tcPr is None:
                    target_cell._element.get_or_add_tcPr()
                for child in cell._element.tcPr:
                    if not child.tag.endswith('gridSpan') and not child.tag.endswith('vMerge'):
                        try:
                            target_cell._element.tcPr.append(deepcopy(child))
                        except:
                            pass
    
    # Copiar propiedades de tabla
    if source_table._element.tblPr is not None:
        if new_table._element.tblPr is None:
            new_table._element.get_or_add_tblPr()
        new_table._element.tblPr.clear()
        for child in source_table._element.tblPr:
            new_table._element.tblPr.append(deepcopy(child))
    
    if source_table._element.tblGrid is not None:
        new_table._element.tblGrid.clear()
        for child in source_table._element.tblGrid:
            new_table._element.tblGrid.append(deepcopy(child))
    
    return new_table


def add_image_to_docx(docx_path, image_path, template_path, step_desc="", test_case_id=""):
    """Agrega una imagen al documento DOCX copiando la estructura exacta del template."""
    docx_file = Path(docx_path)
    
    replacements = {
        "{{TEST_ID}}": test_case_id,
        "{{STEP_DESC}}": step_desc,
    }

    if docx_file.exists():
        document = Document(docx_path)
        template_doc = Document(template_path)

        if template_doc.tables:
            template_table = template_doc.tables[0]
            new_table = _copy_table_style_and_content(template_table, document)
            
            for row in new_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
            
            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture(image_path, width=Inches(5.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document = Document(template_path)
        
        for paragraph in document.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
        
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(image_path, width=Inches(5.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(docx_path)


def add_step_table(docx_path, template_path, step_desc, test_case_id, description_html=""):
    """Agrega tabla de paso + texto descriptivo personalizado."""
    docx_file = Path(docx_path)
    
    replacements = {
        "{{TEST_ID}}": test_case_id,
        "{{STEP_DESC}}": step_desc,
    }

    if docx_file.exists():
        document = Document(docx_path)
        template_doc = Document(template_path)
        
        if template_doc.tables:
            template_table = template_doc.tables[0]
            new_table = _copy_table_style_and_content(template_table, document)
            
            for row in new_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
            
            # Agregar texto descriptivo personalizado del editor
            if description_html and description_html.strip():
                _html_to_docx_paragraphs(description_html, document)
        
        document.save(docx_path)
        
    else:
        document = Document(template_path)
        
        for paragraph in document.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, replacements)
        
        # Agregar texto descriptivo personalizado del editor
        if description_html and description_html.strip():
            _html_to_docx_paragraphs(description_html, document)
        
        document.save(docx_path)


def add_evidence_image(docx_path, image_path):
    """Agrega una imagen de evidencia."""
    docx_file = Path(docx_path)
    
    if not docx_file.exists():
        raise FileNotFoundError(f"El documento {docx_path} no existe.")
    
    document = Document(docx_path)
    
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(5.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.save(docx_path)
